import configuration
import logging as log
from canvasapi import Canvas
import csv
import math
import time
import requests
import re
import unicodedata
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from docx.api import Document
import io
from doc2docx import convert

MIN_SYLLABI_LENGTH = 9000  # The minimum length of a syllabus to be considered posted. 2024SP template was ~7500 characters

config = configuration.Configuration()
config.load_config_file()

ucfcanvas = Canvas("https://champlain.instructure.com", config.api_key)


def get_course_ids(term_id, account_id=3171, has_users=True):
    """Get the course ids from the term id

    Args:
        term_id (int): The term id
        account_id (int): The account id. Defaults to 3171 (Champlain Trad on-campus)
    """
    account = ucfcanvas.get_account(account_id)
    return [
        course.id
        for course in account.get_courses(
            enrollment_term_id=term_id, with_enrollments=has_users
        )
    ]


def termid_from_name(name, root_account_id=283):
    """Get the term id from the term name

    Args:
        name (str): The name of the term E.g. '2023FA'
        root_account_id (int): The root account id. Defaults to 283 (Champlain College)
    """
    for term in ucfcanvas.get_account(root_account_id).get_enrollment_terms():
        if term.name == name:
            return term.id
    return None


def get_course_ids_in_term(term_name: str):
    """Get the courses in a term

    Args:
        term_name (str): The name of the term to get courses for E.g. '2023FA'

    Returns:
        list: A list of course ids
    """
    term_id = termid_from_name(term_name)
    course_ids = get_course_ids(term_id)
    return course_ids


def get_course_facts(course_id, check_for_google_docs=False):
    """Get basic course facts from the course id"""
    course = ucfcanvas.get_course(course_id, include="syllabus_body")

    course_data = {
        "course_name": course.name,
        "course_id": course_id,
        "course_code": course.course_code,
        "faculty_email": (", ").join(
            [teacher.email for teacher in course.get_users(enrollment_type="teacher")]
        ),
        "sis_id": course.sis_course_id,
    }
    if check_for_google_docs:
        if course.syllabus_body:
            course_data["uses_google_docs"] = (
                "https://docs.google.com/document/" in course.syllabus_body
            )
        else:
            course_data["uses_google_docs"] = False

    return course_data


def get_syllabus_html(course_id):
    """Get the syllabus html from the course id"""
    course = ucfcanvas.get_course(course_id, include="syllabus_body")
    syllabus = course.syllabus_body
    if not syllabus:
        return ""

    return syllabus


def is_syllabus_posted(course_id=None, syllabus=None):
    """Check if a syllabus is posted for a course

    Args:
        course_id (int): Optional - The course id
        syllabus (str): Optional - The syllabus text
    Returns:
        bool: True if a syllabus is posted, False otherwise
    """
    if not syllabus and not course_id:
        raise ValueError("Must provide either a course_id or syllabus")

    if not syllabus:
        syllabus = get_syllabus_html(course_id)

    if "instructure_file_link" in syllabus:  # Has a file attached
        return True
    elif len(syllabus) > MIN_SYLLABI_LENGTH:
        return True
    else:
        return False


def find_file_ids(syllabus_text):
    """Find the file links in the syllabus text and return as download links"""
    soup = BeautifulSoup(syllabus_text, "html.parser")
    # All file links will be in an anchor tag with a class of 'instructure_file_link'
    # The href attribute will contain the link to the file
    # Find all anchor tags with the class of 'instructure_file_link'
    # and return the href attribute
    file_links = []
    for link in soup.find_all("a", class_="instructure_file_link"):
        try:
            if link.get("href").split("?")[0].split("/")[-1].isnumeric():
                id = link.get("href").split("?")[0].split("/")[-1]
            else:
                id = link.get("href").split("?")[0].split("/")[-2]
            if not id.isnumeric():
                log.error(
                    f"Error: {id} is not a valid file id\nLink: {link.get('href')}"
                )
                continue
            file_links.append(id)
        except Exception as e:
            log.error(f"Error: {e}\nSkipping: {link.get('href')}")
            continue
    return file_links


def get_text_from_files(syllabus_text):
    """Get the text from the files linked in the syllabus text"""
    file_ids = find_file_ids(syllabus_text)
    text = ""
    if file_ids:
        for file_id in file_ids:
            try:
                file = ucfcanvas.get_file(int(file_id))
            except Exception as e:
                log.error(f"Error getting file: {e}\n File id: {file_id}")
                continue
            for x in range(5):  # Try up to 5 times to get the file
                try:
                    # Check if the file is a pdf and extract the text
                    if file.__getattribute__("content-type") == "application/pdf":
                        contents = file.get_contents(binary=True)
                        on_fly_mem_obj = io.BytesIO(contents)
                        pdf_file = PdfReader(on_fly_mem_obj)
                        for page in pdf_file.pages:
                            text += page.extract_text()

                        # Replace all occurrences of newlines that do not occur after a period with an empty string
                        text = re.sub("(?<!\\.)\n", " ", text)
                        break
                    # Check if the file is a docx and extract the text
                    elif (
                        file.__getattribute__("content-type")
                        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    ):
                        contents = file.get_contents(binary=True)
                        on_fly_mem_obj = io.BytesIO(contents)
                        docx_file = Document(on_fly_mem_obj)
                        text += "\n".join([p.text for p in docx_file.paragraphs])
                    elif file.__getattribute__("content-type") == "application/msword":
                        # An awkward way to convert doc to docx is required as the docx library does not support doc files
                        contents = file.get_contents(binary=True)
                        # Write contents to a doc file in the ./reports folder called "temp.doc"
                        with open("./reports/temp.doc", "wb") as f:
                            f.write(contents)
                        # Convert the doc file to a docx file
                        convert(
                            "./reports/temp.doc", "./reports/temp.docx"
                        )  # convert from doc to docx
                        # read the contents of temp.docx
                        docx_file = Document("./reports/temp.docx")
                        text += "\n".join([p.text for p in docx_file.paragraphs])
                    else:
                        log.error(
                            f"Error: {file.__getattribute__('content-type')} is not a supported file type"
                        )
                    break  # Break out of the for loop, we got the file
                except requests.exceptions.ConnectionError as e:
                    log.error(f"Error getting file: {e}\n File id: {file_id}")
                    continue
    return text
